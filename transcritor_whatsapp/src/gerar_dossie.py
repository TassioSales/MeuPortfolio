# -*- coding: utf-8 -*-
"""Gera um dossie cronologico (.docx) juntando os achados de audio e texto.
Uso: python src/gerar_dossie.py
Saida: dados/saidas/Dossie_Walef.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

AQUI = os.path.dirname(os.path.abspath(__file__))
SAIDAS = os.path.normpath(os.path.join(AQUI, "..", "dados", "saidas"))
DEST = os.path.join(SAIDAS, "Dossie_Walef.docx")

doc = Document()

def h(txt, size=13, cor=(0,0,0), bold=True, space=6):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.bold = bold; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*cor)
    p.space_after = Pt(space)
    return p

def linha(quem, hora, txt, destaque=False):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{hora} — {quem}: ")
    r.bold = True
    r2 = p.add_run(f"“{txt}”")
    if destaque:
        r2.italic = True
        r2.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

# ---- Capa
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("DOSSIÊ — REGISTRO DE OCORRÊNCIAS NO TRABALHO"); r.bold=True; r.font.size=Pt(16)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Conversas de WhatsApp (texto + áudios transcritos) — Grupo “TI Biomundo - Interno”").italic=True

doc.add_paragraph()
h("Partes")
doc.add_paragraph("• Empregado: Tássio Lucian de Jesus Sales (T.I.)")
doc.add_paragraph("• Superior/gestor: Walef (identificado no grupo como “~W. B.”, número +55 61 9657-4839)")
doc.add_paragraph("• Substituto mencionado: Natan (+55 61 9416-3747)")

# ---- Aviso
h("Aviso importante", cor=(0xB0,0,0))
doc.add_paragraph(
    "Este documento é uma ORGANIZAÇÃO CRONOLÓGICA de mensagens, para leitura e para "
    "apresentação a um(a) advogado(a) trabalhista. NÃO é peça jurídica nem laudo. As "
    "transcrições de áudio foram feitas automaticamente e podem conter pequenos erros; "
    "os trechos devem ser conferidos nos arquivos originais (áudio/print) antes de qualquer uso."
)

# ---- Resumo do padrao
h("Resumo do padrão observado")
for b in [
    "Cobrança dura e culpabilização repetida (ex.: “a culpa é exclusivamente sua”).",
    "Ameaças VELADAS de demissão em forma de crítica de desempenho (ex.: “não vai dar certo”, "
    "“tô vendo seu nível cair”, “tá avisado”).",
    "Nos áudios, sinalização explícita de desligamento (11/06/2026).",
    "PONTO CENTRAL: a pressão recai sobre os momentos de SAÚDE do empregado (05/05, 11/06 e 03/07).",
    "Substituto (Natan) sendo treinado e recebendo os acessos do empregado (03/07/2026).",
]:
    doc.add_paragraph(b, style="List Bullet")

# ---- Linha do tempo
h("Linha do tempo (resumo)")
tab = doc.add_table(rows=1, cols=3); tab.style = "Light Grid Accent 1"
hdr = tab.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "Data", "Ocorrência", "Fonte"
for d, o, f in [
    ("06/02/2026", "Culpabilização repetida + cobrança de atraso", "Áudio"),
    ("30/04/2026", "Ameaça velada: “não vai dar certo / tô vendo seu nível cair / tá avisado”", "Texto"),
    ("05/05/2026", "Passou mal (quase desmaiou); cobrado por “autonomia” e setor desfalcado", "Texto"),
    ("11/06/2026", "Desligamento sinalizado por faltas de saúde", "Áudio"),
    ("03/07/2026", "De ATESTADO (ciente); exigido trabalho o dia todo; senha do certificado e “vc paga”; repasse ao Natan", "Texto"),
]:
    c = tab.add_row().cells; c[0].text, c[1].text, c[2].text = d, o, f

# ---- Detalhamento
doc.add_page_break()
h("Detalhamento por data", size=14)

h("06/02/2026 — culpabilização e pressão (áudios)")
linha("Walef","09:27","...foi lá pro pessoal da BLP. Que tu não deu conta. A demanda inicial é tua.", True)
linha("Walef","09:38","larga tudo, mais prioridades... a gente tem que entregar no mês dois, se não entregar vai ficar mal pra nós... sume a bronca e cai pra dentro... é dar ruim.", True)
linha("Walef","09:45","a culpa é especificamente sua... se não der certo a culpa é exclusivamente sua, irmão.", True)
linha("Walef","08:49","45 minutos de atraso... não tem como aceitar esse argumento.")

h("30/04/2026 — ameaça velada de demissão (texto)")
linha("Walef","15:13","tem coisa que eu não vou conseguir ficar te pedindo, te mostrando, tu é pleno.")
linha("Walef","15:14","se não progredir não vai dar certo velho, sinto muito, tô vendo seu nível cair.", True)
linha("Walef","15:15","tá avisado, preciso que vc se organize melhor e seja mais responsável.", True)
linha("Walef","15:15","infelizmente é nosso cenário velho, não consigo fazer muito.", True)
linha("Walef","15:19","seu número de atendimento no seu período aqui é muito pequeno, fui mostrar para diretoria, o número ficou muito baixo.")

h("05/05/2026 — episódio de saúde (texto)")
linha("Tássio","09:38","passei mau aqui, tô indo no Amor e Saúde do Cruzeiro pra consultar.")
linha("Walef","12:35","seu gestor sou eu, você não deve se dirigir diretamente assim sem alinhar comigo antes.", True)
linha("Walef","12:35","mais uma vez tendo que te reportar sobre autonomia.", True)
linha("Walef","12:37","setor ficou desfalcado e recebi reclamação, nem sabia de nada.", True)
linha("Tássio","12:41","minha glicemia deu baixa. A médica me deu um atestado de dois dias.")

h("11/06/2026 — desligamento por faltas de saúde (áudios)")
linha("Walef","07:06","tá foda essas duas falta... nosso setor é de sustentação, um dia que tu falta já desanda tudo... não dá pra ficar dessa forma.", True)
linha("Walef","07:08","sei que é uma questão de saúde... precisa de pessoa com disponibilidade... se tu não é esse cara, infelizmente eu não consigo segurar... acredito que já sabe o que a gente vai chegar.", True)

h("03/07/2026 — de ATESTADO, mas exigido trabalho o dia todo (texto)")
doc.add_paragraph(
    "Fato informado pelo empregado: neste dia estava de ATESTADO. Comunicou e ENTREGOU O "
    "ATESTADO EM MÃOS ao Walef (empregador ciente). Ainda assim, foi exigido trabalho durante "
    "todo o expediente (relatório de 142 bases, emissão de certificados digitais, repasse de "
    "acessos e treinamento do substituto Natan)."
)
linha("Walef","15:36","123456* — A PADRÃO É ESSA / temos um padrão para todos tu sabe / não vou mudar o padrão.", True)
linha("Walef","15:39","mas foi erro seu, podia ter consultado antes; tu faz isso quase todo mês.")
linha("Walef","15:40","se tiver custo vai ter pagar aí.", True)
linha("Tássio","15:41","uai, se tá falando que vou ter que pagar um certificado.")
linha("Walef","15:43","chama o Natan e passa o processo pra ele fazer o outro.")
doc.add_paragraph(
    "Observações: (1) o “padrão” de senha dos certificados digitais da empresa é “123456*”, "
    "usado para todos — falha grave de segurança em credencial que assina/emite em nome da "
    "empresa. (2) Repassar ao empregado o custo de reemissão por um erro de digitação é indevido: "
    "o empregador arca com os riscos do negócio (CLT, art. 2º) e o desconto salarial é restrito "
    "(CLT, art. 462). O Walef recuou ao ser questionado."
)

# ---- Provas a guardar
doc.add_page_break()
h("O que preservar (checklist de provas)")
for b in [
    "Os arquivos de áudio originais (.opus) — não apagar do celular.",
    "Prints das conversas de texto mostrando data, hora e nome/numero do remetente.",
    "Cópia do ATESTADO de 03/07/2026 e comprovante/protocolo de entrega ao RH; nomes de testemunhas da entrega em mãos.",
    "Atestados anteriores (ex.: 05/05/2026) e registros de idas ao médico.",
    "E-mails de “ponto de controle” e demais mensagens citadas.",
]:
    doc.add_paragraph(b, style="List Bullet")

h("Recomendação")
doc.add_paragraph(
    "Procurar um(a) advogado(a) trabalhista, apresentando este material. Os pontos de maior peso "
    "são: (a) a pressão e o encaminhamento de desligamento ligados a ausências de saúde "
    "(possível dispensa discriminatória — Súmula 443 do TST); e (b) a exigência de trabalho "
    "durante atestado, com o empregador ciente."
)

doc.save(DEST)
print("Dossie gerado em:", DEST)
